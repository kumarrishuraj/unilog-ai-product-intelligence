"""
Robust ingestion for the Unilog reference workbooks.

The brief warns these files may contain merged cells, multi-row headers, side-by-side
tables, notes rows and empty columns.  Nothing here assumes a sheet name, a header
row index or a column position: every sheet is *sniffed*.

Design contract
---------------
Each ``load_*`` function returns ``(records, source_note)`` and returns an empty
list when the file is absent.  The caller then falls back to the bootstrap miner.
This is what lets the whole system run with or without the official reference pack.
"""
from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

try:                                    # openpyxl is optional at import time
    from openpyxl import load_workbook
    _HAVE_XLSX = True
except Exception:                       # pragma: no cover
    _HAVE_XLSX = False

try:
    import docx                         # python-docx
    _HAVE_DOCX = True
except Exception:                       # pragma: no cover
    _HAVE_DOCX = False

from backend.normalization.text import normalize_whitespace, repair_mojibake


# ---------------------------------------------------------------------------
# Generic sheet reading
# ---------------------------------------------------------------------------
def _cell(value) -> str:
    if value is None:
        return ""
    return normalize_whitespace(repair_mojibake(str(value)))


def read_sheet_matrix(path: Path, sheet_name: str) -> List[List[str]]:
    """
    Read a worksheet into a dense string matrix with merged cells un-merged
    (the top-left value is propagated across the merged range).
    """
    wb = load_workbook(path, read_only=False, data_only=True)
    ws = wb[sheet_name]
    matrix = [[_cell(c) for c in row] for row in ws.iter_rows(values_only=True)]

    for rng in list(getattr(ws, "merged_cells", []).ranges if hasattr(ws, "merged_cells") else []):
        r0, c0, r1, c1 = rng.min_row - 1, rng.min_col - 1, rng.max_row - 1, rng.max_col - 1
        if r0 >= len(matrix) or c0 >= len(matrix[r0]):
            continue
        anchor = matrix[r0][c0]
        for r in range(r0, min(r1 + 1, len(matrix))):
            for c in range(c0, min(c1 + 1, len(matrix[r]))):
                if not matrix[r][c]:
                    matrix[r][c] = anchor
    wb.close()
    return matrix


def sheet_names(path: Path) -> List[str]:
    if not _HAVE_XLSX or not path.exists():
        return []
    wb = load_workbook(path, read_only=True)
    names = list(wb.sheetnames)
    wb.close()
    return names


def _score_header_row(row: Sequence[str], below: Sequence[Sequence[str]]) -> float:
    """
    Heuristic 'is this the header row?' score: mostly-populated, mostly-textual,
    distinct labels, and followed by rows of similar width.
    """
    cells = [c for c in row if c]
    if len(cells) < 2:
        return 0.0
    distinct = len({c.lower() for c in cells}) / len(cells)
    textual = sum(1 for c in cells if not re.fullmatch(r"[\d.,%$/-]+", c)) / len(cells)
    fill = len(cells) / max(1, len(row))
    # Header rows are usually short strings, not sentences.
    brevity = sum(1 for c in cells if len(c) <= 60) / len(cells)
    support = 0.0
    if below:
        widths = [len([c for c in r if c]) for r in below[:5]]
        if widths:
            support = min(1.0, (sum(widths) / len(widths)) / max(1, len(cells)))
    return 0.30 * distinct + 0.25 * textual + 0.15 * fill + 0.10 * brevity + 0.20 * support


def sniff_table(matrix: List[List[str]], max_scan: int = 25) -> Tuple[int, List[str]]:
    """
    Find the header row index and its labels.  Handles leading title/notes rows and
    two-row headers (a blank upper cell inherits from the row above).
    """
    best_idx, best_score = 0, -1.0
    for i in range(min(max_scan, len(matrix))):
        score = _score_header_row(matrix[i], matrix[i + 1:i + 6])
        if score > best_score:
            best_idx, best_score = i, score

    header = list(matrix[best_idx]) if best_idx < len(matrix) else []

    # Two-row header: if the next row is also header-ish and this one has gaps,
    # concatenate them ('Attribute' / 'Label' -> 'Attribute Label').
    if best_idx + 1 < len(matrix):
        nxt = matrix[best_idx + 1]
        gaps = sum(1 for c in header if not c)
        if gaps and _score_header_row(nxt, matrix[best_idx + 2:best_idx + 7]) > 0.45:
            merged, carry = [], ""
            for j in range(max(len(header), len(nxt))):
                top = header[j] if j < len(header) else ""
                bot = nxt[j] if j < len(nxt) else ""
                carry = top or carry
                merged.append(normalize_whitespace(f"{carry} {bot}").strip() if bot else carry)
            if sum(1 for c in merged if c) > sum(1 for c in header if c):
                header = merged
                best_idx += 1

    # De-duplicate blank/repeated labels so dict rows do not collide.
    seen: Dict[str, int] = {}
    out: List[str] = []
    for j, h in enumerate(header):
        name = h or f"column_{j + 1}"
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 0
        out.append(name)
    return best_idx, out


def sheet_records(path: Path, sheet_name: str) -> List[Dict[str, str]]:
    """Read a sheet into dicts, auto-detecting the header row. Drops empty rows."""
    matrix = read_sheet_matrix(path, sheet_name)
    if not matrix:
        return []
    hdr_idx, header = sniff_table(matrix)
    records: List[Dict[str, str]] = []
    for row in matrix[hdr_idx + 1:]:
        if not any(row):
            continue
        rec = {header[j]: (row[j] if j < len(row) else "") for j in range(len(header))}
        if any(v for v in rec.values()):
            records.append(rec)
    return records


def workbook_records(path: Path) -> Dict[str, List[Dict[str, str]]]:
    """Every sheet of a workbook, sniffed independently."""
    if not _HAVE_XLSX or not path.exists():
        return {}
    out: Dict[str, List[Dict[str, str]]] = {}
    for name in sheet_names(path):
        try:
            out[name] = sheet_records(path, name)
        except Exception:                # a broken sheet must not kill the load
            out[name] = []
    return out


def csv_records(path: Path) -> List[Dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as fh:
        return [{k: _cell(v) for k, v in row.items()} for row in csv.DictReader(fh)]


# ---------------------------------------------------------------------------
# Column matching
# ---------------------------------------------------------------------------
def _norm_col(name: str) -> str:
    return re.sub(r"[^a-z0-9]", "", str(name or "").lower())


def find_column(header: Iterable[str], *candidates: str) -> Optional[str]:
    """
    Locate a column by fuzzy name, tolerating spacing/underscore/case differences.
    Exact-normalised match first, then substring, so 'MANUFACTURER_NAME' matches
    'Manufacturer Name' and 'Mfr Name' matches 'name'.
    """
    cols = list(header)
    normed = {c: _norm_col(c) for c in cols}
    for cand in candidates:
        target = _norm_col(cand)
        for c in cols:
            if normed[c] == target:
                return c
    for cand in candidates:
        target = _norm_col(cand)
        for c in cols:
            if target and target in normed[c]:
                return c
    return None


def pick_sheet(book: Dict[str, List[Dict[str, str]]], *keywords: str) -> Optional[str]:
    """Choose the sheet whose name best matches the keywords, else the largest."""
    if not book:
        return None
    for kw in keywords:
        k = _norm_col(kw)
        for name in book:
            if k and k in _norm_col(name):
                return name
    return max(book, key=lambda n: len(book[n])) if book else None


# ---------------------------------------------------------------------------
# DOCX (content guidelines)
# ---------------------------------------------------------------------------
def read_docx_blocks(path: Path) -> List[Dict[str, str]]:
    """
    Flatten a .docx into retrievable blocks: {'heading': ..., 'text': ...}.
    Used to build the content-guidelines RAG collection.
    """
    if not _HAVE_DOCX or not path.exists():
        return []
    doc = docx.Document(str(path))
    blocks: List[Dict[str, str]] = []
    heading = ""
    buf: List[str] = []

    def flush():
        if buf:
            blocks.append({"heading": heading, "text": normalize_whitespace(" ".join(buf))})
            buf.clear()

    for para in doc.paragraphs:
        text = normalize_whitespace(para.text)
        if not text:
            continue
        if (para.style.name or "").lower().startswith("heading"):
            flush()
            heading = text
        else:
            buf.append(text)
    flush()

    for t_i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [normalize_whitespace(c.text) for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append({"heading": f"{heading} (table {t_i + 1})", "text": "\n".join(rows)})
    return blocks
